import streamlit as st
import re
import io
import zipfile
import os
from docx import Document

# --- 設定: 除外したい単語リスト ---
IGNORE_LIST = [
    '参加者', '話者', '詳細', 'まとめ', '日時', 'Source', 'source', '文字起こし', 'メモ', '長さ', 'Time', 'Unknown',
    'ENG', 'JPN', 'ENG/JPN', 'ENG_JPN', 'JST', 'Gemini', 'によるメモ', 'のコピー', '標準', 'インタビュー', '対象者',
    '会議の録画', '招待済み', '添付ファイル', 'mp4', 'm4a', 'wav', 'docx', 'txt', 'pdf', 'com', 'jp', 'ac',
    'Speaker', '筑波大学'
]

def is_valid_name(name):
    """名前として適切か判定する"""
    clean_name = name.strip()
    if not clean_name:
        return False
    if len(clean_name) <= 1:
        return False
    if clean_name.isdigit(): 
        return False
    
    # 除外リストに含まれるかチェック（大文字小文字無視）
    for ignore in IGNORE_LIST:
        # 完全一致チェック
        if ignore.lower() == clean_name.lower():
            return False
        # Speaker_A のような既存の置換ネームも除外
        if "speaker" in clean_name.lower():
            return False
            
        # 日付形式の除外 (数字と記号の混在)
        if re.search(r'\d', clean_name) and re.search(r'[\/\-_]', clean_name):
            if '@' not in clean_name: # メアドは許可
                return False
    return True

def extract_names(text, filename=""):
    """テキストとファイル名から名前・メールアドレス候補をすべて抽出する"""
    potential_names = set()

    # 1. イニシャル付きの名前パターン (最優先追加)
    # 例: R.Okuzumi, X.Su, H.Sakai
    # [大文字1文字] [ドット] [大文字] [英字1文字以上]
    pattern_initial = r'\b[A-Z]\.[A-Z][a-zA-Z]+'
    matches_initial_text = re.findall(pattern_initial, text)
    matches_initial_file = re.findall(pattern_initial, filename)
    potential_names.update(matches_initial_text)
    potential_names.update(matches_initial_file)

    # 2. メールアドレス
    pattern_email = r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'
    matches_email_text = re.findall(pattern_email, text)
    matches_email_file = re.findall(pattern_email, filename)
    potential_names.update(matches_email_text)
    potential_names.update(matches_email_file)

    # 3. 本文中の '名前: ' パターン
    pattern_colon = r'(?:^|\n)(?:\[.*?\]\s*)?([^\n\r：:]{2,20}?)\s*[:：]'
    matches_colon = re.findall(pattern_colon, text)
    potential_names.update(matches_colon)

    # 4. ファイル名やヘッダーにある括弧内の文字列
    # 括弧の中身を取り出し、さらに " - " 等で分割して評価する
    base_name = os.path.splitext(filename)[0]
    search_target = base_name + "\n" + text[:500] 
    pattern_bracket = r'[（\(]([^）\)\n\r]+?)[）\)]'
    matches_bracket = re.findall(pattern_bracket, search_target)
    
    for content in matches_bracket:
        # 括弧の中身を区切り文字で分割してみる (例: "Speaker_C - R.Okuzumi")
        parts = re.split(r'[\s\-_/]+', content)
        # 分割前の全体も候補に入れる
        potential_names.add(content)
        # 分割後のパーツも候補に入れる
        for p in parts:
            potential_names.add(p)

    # フィルタリング
    unique_names = set()
    for name in potential_names:
        # 記号を除去して純粋な名前部分だけでチェック
        clean = name.strip(" -_")
        if is_valid_name(clean):
            unique_names.add(clean)
    
    # 名前が長い順にソート
    return sorted(list(unique_names), key=len, reverse=True)

def generate_name_map(names):
    """名前リストから置換マップ(Speaker_A...)を作成"""
    name_map = {}
    chars = "ABCDEFGHIJKLMNOPQRSTUVWXYZ"
    for i, name in enumerate(names):
        replacement = f"Speaker_{chars[i % len(chars)]}"
        if i >= len(chars):
            replacement += str(i)
        name_map[name] = replacement
    return name_map

def process_content(content, filename):
    names = extract_names(content, filename)
    name_map = generate_name_map(names)

    # 本文の置換
    new_content = content
    for original, new in name_map.items():
        new_content = new_content.replace(original, new)

    # ファイル名の置換
    name_part, ext = os.path.splitext(filename)
    new_name_part = name_part
    for original, new in name_map.items():
        if original in new_name_part:
            new_name_part = new_name_part.replace(original, new)
    
    new_filename = new_name_part + ext
    return new_filename, new_content, name_map

# --- ファイル処理ラッパー ---
def process_text_file(file_obj):
    try:
        bytes_data = file_obj.getvalue()
        try:
            content = bytes_data.decode('utf-8')
        except UnicodeDecodeError:
            content = bytes_data.decode('cp932', errors='ignore')
    except:
        return None, None
    
    new_filename, new_content, _ = process_content(content, file_obj.name)
    return new_filename, new_content.encode('utf-8')

def process_docx_file(file_obj):
    try:
        doc = Document(file_obj)
    except:
        return None, None

    # 全文取得
    full_text_list = []
    for para in doc.paragraphs:
        full_text_list.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text_list.append(cell.text)
    
    full_text_joined = "\n".join(full_text_list)
    
    names = extract_names(full_text_joined, file_obj.name)
    name_map = generate_name_map(names)

    # 置換実行
    for para in doc.paragraphs:
        for original, new in name_map.items():
            if original in para.text:
                para.text = para.text.replace(original, new)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for original, new in name_map.items():
                    if original in cell.text:
                        cell.text = cell.text.replace(original, new)

    name_part, ext = os.path.splitext(file_obj.name)
    new_name_part = name_part
    for original, new in name_map.items():
        if original in new_name_part:
            new_name_part = new_name_part.replace(original, new)
    new_filename = new_name_part + ext

    output_stream = io.BytesIO()
    doc.save(output_stream)
    return new_filename, output_stream.getvalue()

# --- アプリ画面 ---
st.title("🕵️ 文字起こし匿名化ツール v4")
st.markdown("""
以下の情報を一括で `Speaker_X` 等に変換します。
* **名前**（会話の「名前:」）
* **英字氏名**（`R.Okuzumi`, `X.Su` など）
* **メールアドレス**
* **ファイル名の括弧内の氏名**

対応形式: `.txt`, `.md`, `.csv`, `.docx`
""")

uploaded_files = st.file_uploader("ファイルをドラッグ＆ドロップ", accept_multiple_files=True)

if uploaded_files:
    if st.button(f"{len(uploaded_files)} ファイルを処理開始"):
        progress_bar = st.progress(0)
        zip_buffer = io.BytesIO()
        processed_count = 0
        
        with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zip_file:
            for i, file_obj in enumerate(uploaded_files):
                filename = file_obj.name
                ext = os.path.splitext(filename)[1].lower()
                
                if ext == '.docx':
                    new_name, new_data = process_docx_file(file_obj)
                else:
                    new_name, new_data = process_text_file(file_obj)
                
                if new_name and new_data:
                    zip_file.writestr(new_name, new_data)
                    processed_count += 1
                
                progress_bar.progress((i + 1) / len(uploaded_files))
        
        st.success(f"完了！ {processed_count} / {len(uploaded_files)} ファイル処理済み")
        st.download_button("📦 ZIPをダウンロード", zip_buffer.getvalue(), "anonymized_v4.zip", "application/zip")
